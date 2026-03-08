import logging
from dataclasses import dataclass
from pathlib import Path

logger = logging.getLogger(__name__)

# Максимум символов передаём в модель — ~6000 токенов
_MAX_CHARS = 24_000
# Поддерживаемые расширения
SUPPORTED_EXTENSIONS = {".pdf", ".doc", ".docx", ".txt"}


@dataclass(frozen=True)
class DocumentResult:
    """Результат извлечения текста из документа."""
    text: str
    pages: int        # количество страниц (для PDF) или 0
    truncated: bool   # был ли текст обрезан


def extract_text(file_path: Path) -> DocumentResult:
    """
    Извлекает текст из документа.
    Поддерживает PDF, DOCX, DOC, TXT.
    Возвращает DocumentResult или бросает ValueError при неподдерживаемом формате.
    """
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf(file_path)
    elif suffix in (".docx", ".doc"):
        return _extract_docx(file_path)
    elif suffix == ".txt":
        return _extract_txt(file_path)
    else:
        raise ValueError(f"Неподдерживаемый формат: {suffix}")


def _truncate(text: str) -> tuple[str, bool]:
    """Обрезает текст до лимита. Возвращает (текст, был_обрезан)."""
    text = text.strip()
    if len(text) <= _MAX_CHARS:
        return text, False
    return text[:_MAX_CHARS], True


def _extract_pdf(path: Path) -> DocumentResult:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise RuntimeError("PyMuPDF не установлен. Добавь PyMuPDF в requirements.txt")

    doc = fitz.open(str(path))
    pages = doc.page_count
    parts: list[str] = []

    for page in doc:
        page_text = page.get_text().strip()
        if page_text:
            parts.append(page_text)

    doc.close()
    raw = "\n\n".join(parts)
    text, truncated = _truncate(raw)

    logger.info("📄 PDF: %d стр., %d символов%s", pages, len(text),
                " (обрезан)" if truncated else "")
    return DocumentResult(text=text, pages=pages, truncated=truncated)


def _extract_docx(path: Path) -> DocumentResult:
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx не установлен. Добавь python-docx в requirements.txt")

    doc = Document(str(path))
    parts: list[str] = []

    for para in doc.paragraphs:
        stripped = para.text.strip()
        if stripped:
            parts.append(stripped)

    # Извлекаем текст из таблиц
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)

    raw = "\n".join(parts)
    text, truncated = _truncate(raw)

    logger.info("📝 DOCX: %d символов%s", len(text), " (обрезан)" if truncated else "")
    return DocumentResult(text=text, pages=0, truncated=truncated)


def _extract_txt(path: Path) -> DocumentResult:
    raw = path.read_text(encoding="utf-8", errors="replace").strip()
    text, truncated = _truncate(raw)
    logger.info("📃 TXT: %d символов%s", len(text), " (обрезан)" if truncated else "")
    return DocumentResult(text=text, pages=0, truncated=truncated)
